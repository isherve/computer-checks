"""Build Chapter Four Word document with all figures inserted."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).resolve().parent
DOC_PATH = OUT_DIR.parent / "CHAPTER_FOUR_System_Analysis_and_Design.docx"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading_styled(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.clear()
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True)
    return p


def add_figure(doc, filename, caption, width=5.8):
    path = OUT_DIR / filename
    if not path.exists():
        add_body(doc, f"[Missing figure file: {filename}]", first_line_indent=False)
        add_caption(doc, caption)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=10, bold=True)
        set_cell_shading(hdr[i], "1F4E79")
        run.font.color.rgb = RGBColor(255, 255, 255)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CHAPTER FOUR: SYSTEM ANALYSIS AND DESIGN")
    set_run_font(r, size=16, bold=True)

    # 4.1
    add_heading_styled(doc, "4.1 Introduction", 2)
    add_body(
        doc,
        "This chapter presents the analysis and design of the proposed System for Personal "
        "Computer Checks Using QR Codes, developed for UTB Rubavu Campus. It discusses the "
        "analysis of the existing manual gate-recording system, the requirements identified "
        "for the new system, and the design of the proposed solution. The chapter also "
        "includes system modelling tools such as Use Case Diagrams, Data Flow Diagrams (DFDs), "
        "a Data Dictionary, Entity Relationship Diagrams (ERDs), database normalization, and "
        "the overall system architecture.",
    )

    # 4.2
    add_heading_styled(doc, "4.2 Data analysis and presentation", 2)
    add_body(
        doc,
        "This section presents and analyses information gathered through observation of, and "
        "documentation review at, the UTB Rubavu Campus gate. The areas examined included the "
        "availability of records on laptops and personal computers entering and leaving the "
        "institution, the tools used to capture this information, the process followed by gate "
        "staff when verifying a returning device, and the challenges encountered in retrieving "
        "previously recorded data. The presentation methods used in this analysis include "
        "descriptive tables, process narratives, and diagrams (use case diagrams and flowcharts) "
        "illustrating how the current gate-checking process operates.",
    )

    # 4.3
    add_heading_styled(doc, "4.3 Interpretation of findings/results", 2)
    add_body(
        doc,
        "The analysis of the existing gate-checking process revealed that device records are "
        "captured on paper in notebooks, which are frequently kept apart and sometimes managed "
        "by different personnel. This makes it difficult to locate a specific record when a "
        "device is due to exit, increases the risk of losing or damaging data, and leaves the "
        "process prone to human error and duplication. The findings indicate a clear need for "
        "a centralized, electronic system that can register devices at entry, generate a unique "
        "identifier for each device, and allow gate staff to verify a device quickly and "
        "accurately at exit using that identifier.",
    )

    # 4.4
    add_heading_styled(doc, "4.4 Summary of proposed system", 2)
    add_body(
        doc,
        "The proposed System for Personal Computer Checks Using QR Codes was designed to replace "
        "the manual, paper-based recording process at UTB Rubavu Campus with an electronic, "
        "QR-code-based verification process. The system was implemented using HTML, Bootstrap, "
        "JavaScript, and PHP (with PDO) as the main development technologies, XAMPP as the local "
        "server environment, and MySQL as the database management system for data storage and "
        "management.",
    )
    add_body(
        doc,
        "The study resulted in the development of a system capable of registering personal "
        "computers at the point of entry, generating a scannable QR code for each registered "
        "device, and enabling gate staff to verify a device's identity at exit by scanning its "
        "QR code against the stored record.",
    )

    # 4.5
    add_heading_styled(doc, "4.5 Description of existing system", 2)
    add_body(
        doc,
        "The current system used at UTB Rubavu Campus is a manual, paper-based process in which "
        "notebooks are used to keep track of laptops and personal computers entering the "
        "institution. As described in the project background, this approach is outdated and "
        "exposes recorded data to risk of damage and loss. Records are also not easily kept or "
        "retrieved without consuming significant physical storage space, and the process is "
        "prone to errors and duplication.",
    )
    add_body(
        doc,
        "When a student or guest returns to exit the campus, the gate worker must go back "
        "through the paper records to check whether the device being taken out matches one that "
        "was recorded on entry. Figure 1 below shows the use case diagram of the existing system, "
        "illustrating how students, guests, and gate workers interact with the paper-based process.",
    )
    add_figure(doc, "Figure1_Existing_UseCase.png", "Figure 1: Use Case diagram of the existing system", width=5.5)

    add_body(
        doc,
        "Figure 2 presents a flowchart of the existing system. A flowchart provides a step-by-step "
        "visual representation of a process or workflow; here it outlines the major steps involved "
        "in the current manual gate-checking operation, from recording a device on entry to "
        "verifying and releasing it on exit.",
        first_line_indent=False,
    )
    add_figure(doc, "Figure2_Existing_Flowchart.png", "Figure 2: Flow chart of the existing system", width=4.2)

    # 4.5.1
    add_heading_styled(doc, "4.5.1 Problem with the existing system", 3)
    add_body(
        doc,
        "The analysis of the existing system revealed several challenges that affect effective "
        "device tracking at the campus gate:",
        first_line_indent=False,
    )
    for b in [
        "Finding a previously registered device and its owner among scattered notebooks is difficult and time-consuming.",
        "Recorded data is easily lost because entries are kept in different books, often managed apart and by different people.",
        "The unorganized, mixed record-keeping process increases the likelihood of unexpected errors and inconsistencies.",
        "Paper records are vulnerable to physical damage and take up significant storage space.",
        "Verifying a device at exit relies entirely on manual cross-checking, which is slow and error-prone.",
    ]:
        add_bullet(doc, b)
    add_body(
        doc,
        "These problems reduce the reliability of gate security checks and negatively affect the "
        "efficiency of the entry and exit process for students, staff, and guests.",
        first_line_indent=False,
    )

    # 4.6
    add_heading_styled(doc, "4.6 Description of the new system", 2)
    add_body(
        doc,
        "The proposed system is a web-based System for Personal Computer Checks Using QR Codes, "
        "designed to centralize the recording and verification of personal computers entering and "
        "leaving UTB Rubavu Campus. The system provides gate staff and administrators with a "
        "single point of access to device records, allowing devices to be registered on entry, "
        "matched against a generated QR code, and verified on exit by scanning that code.",
    )
    add_body(
        doc,
        "The system allows authorized users to register a device and its owner's details, generate "
        "a unique QR code for that device, scan the QR code at exit to confirm a match, and "
        "maintain a log of all actions performed for accountability and reporting purposes.",
    )

    add_heading_styled(doc, "4.6.1 Modules description", 3)
    add_body(doc, "The proposed system consists of the following modules:", first_line_indent=False)

    add_heading_styled(doc, "4.6.1.1 User Management Module", 4)
    add_body(
        doc,
        "This module manages user accounts and access control, with the following functions:",
        first_line_indent=False,
    )
    for b in [
        "User login and logout",
        "Password management",
        "User profile management",
        "Role assignment (Administrator, Gate Officer)",
    ]:
        add_bullet(doc, b)
    add_body(
        doc,
        "This ensures secure access to the system and allows different users to perform only the "
        "activities they are authorized to carry out.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "4.6.1.2 Computer Registration Module", 4)
    add_body(
        doc,
        "This module manages the registration of personal computers and their owners as they enter "
        "the institution. It includes:",
        first_line_indent=False,
    )
    for b in [
        "Register a computer's serial number and model",
        "Capture owner's name and contact/owner number",
        "Record the date the device was registered",
        "Edit or update a device record",
    ]:
        add_bullet(doc, b)
    add_body(
        doc,
        "This module provides an accurate, centralized record of every device entering the campus, "
        "replacing the paper notebooks used in the existing system.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "4.6.1.3 QR Code Generation and Verification Module", 4)
    add_body(
        doc,
        "This module generates and verifies the QR code associated with each registered device. "
        "Its functions include:",
        first_line_indent=False,
    )
    for b in [
        "Generate a unique QR code linked to a device's record on registration",
        "Print the generated QR code for the device or its owner",
        "Scan a QR code at the exit gate using a smartphone camera or scanner app",
        "Match scanned data against the stored device record to confirm identity",
    ]:
        add_bullet(doc, b)
    add_body(
        doc,
        "This module is central to the system, as it replaces manual cross-checking with an "
        "automated, reliable verification step at the gate.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "4.6.1.4 Logs and Reporting Module", 4)
    add_body(
        doc,
        "This module records every action performed on a device record and supports reporting for "
        "management purposes. Its functions include:",
        first_line_indent=False,
    )
    for b in [
        "Log each registration, verification, and exit action with a timestamp",
        "Record comments or remarks on an action",
        "Generate reports on devices registered, verified, or flagged over a given period",
    ]:
        add_bullet(doc, b)
    add_body(
        doc,
        "This module supports accountability at the gate and provides administrators with insight "
        "into device movement and system usage.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "4.6.2 System configurations and technology", 3)
    add_body(
        doc,
        "This section outlines the hardware and software requirements needed to develop and operate "
        "the System for Personal Computer Checks Using QR Codes.",
        first_line_indent=False,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Table 1: Hardware specification")
    set_run_font(r, size=11, bold=True)
    add_table(
        doc,
        ["Item", "Requirement"],
        [
            ["Computer", "Required to run the system at the gate"],
            ["Printer", "Required for printing the generated QR codes"],
            ["Smartphone", "Required for scanning QR codes, using a QR-code scanner app or a built-in camera"],
        ],
        col_widths=[1.5, 5.0],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Table 2: Software specification")
    set_run_font(r, size=11, bold=True)
    add_table(
        doc,
        ["Software", "Purpose"],
        [
            ["MySQL", "Free, open-source Relational Database Management System (RDBMS) using SQL, used to insert, select, delete, and update captured information"],
            ["XAMPP", "Free, open-source cross-platform web server solution used to host the system locally"],
            ["PHP (with PDO)", "Backend development"],
            ["JavaScript", "Client-side interaction on the backend/frontend"],
            ["HTML", "Defines the content/structure of the system's pages"],
            ["Bootstrap", "CSS framework responsible for responsive design and browser compatibility on the front end"],
        ],
        col_widths=[1.8, 4.7],
    )

    # 4.7
    add_heading_styled(doc, "4.7 Illustration of new system", 2)
    add_body(
        doc,
        "This project focused on understanding what the System for Personal Computer Checks Using "
        "QR Codes must do and how it should solve the problems identified in the existing manual "
        "gate-checking process. At this stage, emphasis was placed on identifying user needs, "
        "system inputs and outputs, data flow, and overall system behaviour.",
    )
    add_body(
        doc,
        "System design follows system analysis and explains how the proposed system was built to "
        "meet the identified requirements. This includes defining the structure of the system, "
        "breaking it into smaller components, and deciding how these components interact with one "
        "another. The design phase translates the system requirements into a form that can be "
        "implemented using the chosen technologies.",
    )

    add_heading_styled(doc, "4.7.1 Functional Diagram", 3)
    add_body(
        doc,
        "A functional diagram is used to represent the main functions of the System for Personal "
        "Computer Checks Using QR Codes and how users interact with them. It shows the sequence of "
        "actions performed by the system and the relationships between different system functions.",
    )
    add_body(
        doc,
        "In this project, the functional diagram illustrates how gate officers and administrators "
        "interact with the platform. Gate officers register devices on entry, generate and print "
        "QR codes, and scan codes at exit to verify a match. Administrators manage user accounts, "
        "review logs, and generate reports on device activity.",
    )
    add_figure(doc, "Figure3_Functional_Diagram.png", "Figure 3: Functional Diagram", width=5.8)

    add_heading_styled(doc, "4.7.1.1 Data Flow Diagram", 4)
    add_body(
        doc,
        "A Data Flow Diagram (DFD) is a graphical representation of how data moves through a "
        "system. It shows where data comes from, how it is processed, where it is stored, and "
        "where it goes. It does not represent control flow, loops, or decision rules; those are "
        "better captured in a flowchart. It is a useful tool for communicating with users, "
        "managers, and other personnel, and for analysing both existing and proposed systems. [8]",
    )

    add_heading_styled(doc, "4.7.1.1.1 DFD Level 0", 4)
    add_figure(doc, "Figure4_DFD_Level0.png", "Figure 4: DFD Level 0", width=6.0)
    add_body(
        doc,
        "The diagram above represents the Data Flow Diagram (DFD) Level 0 for the System for "
        "Personal Computer Checks Using QR Codes. It provides a high-level view of how data flows "
        "within the system and how different users interact with it. At this level, the entire "
        "system is represented as a single process.",
        first_line_indent=False,
    )
    add_body(
        doc,
        "There are two main external entities involved in the system: the Gate Officer and the "
        "Administrator. Gate officers interact with the system to register devices, generate QR "
        "codes, and verify devices at exit. Administrators interact with the system to manage "
        "users, review device logs, and generate reports.",
    )
    add_body(
        doc,
        "The data flow in this DFD illustrates the exchange of information between users and the "
        "main system process, including device registration details, generated QR-code data, "
        "scan/verification requests, and the responses generated by the system. The DFD Level 0 "
        "helps in understanding the overall system boundaries and the general flow of data from "
        "input to output.",
    )

    add_heading_styled(doc, "4.7.1.1.2 DFD Level 1", 4)
    add_body(
        doc,
        "A DFD Level 1 provides a more detailed view of how data flows within a system by expanding "
        "the single process shown in DFD Level 0 into multiple sub-processes, data stores, and "
        "data flows.",
    )
    add_body(
        doc,
        "In the System for Personal Computer Checks Using QR Codes, the DFD Level 1 breaks the "
        "system down into core processes such as registering a device, generating a QR code, "
        "scanning and verifying a device at exit, logging actions, and administrator authentication. "
        "The DFD Level 1 consists of four main components:",
    )
    for b in [
        "External Entities: users or systems outside the platform that interact with it — in this system, Gate Officers and Administrators.",
        "Processes: system activities that transform input data into output, such as registering a device, generating a QR code, and verifying a scanned code.",
        "Data Flows: show how data moves between external entities, processes, and data stores in a specific direction.",
        "Data Stores: locations where data is stored for later use, such as the Computer_info, Users, and Logs tables.",
    ]:
        add_bullet(doc, b)

    add_figure(doc, "Figure5_DFD_Level1.png", "Figure 5: DFD Level 1", width=6.2)
    add_body(
        doc,
        "The diagram above illustrates how data will flow throughout the System for Personal "
        "Computer Checks Using QR Codes, from device registration through to verification at exit, "
        "and how gate officers and administrators interact with the system's internal processes.",
        first_line_indent=False,
    )
    add_body(
        doc,
        "Additionally, Figure 6 below is a flowchart which illustrates how UTB Rubavu Campus will "
        "use the System for Personal Computer Checks Using QR Codes in operation. It shows clearly "
        "how information flows within the system, and how users and administrators interact with it.",
    )
    add_figure(doc, "Figure6_New_System_Flowchart.png", "Figure 6: Flow chart of the new system", width=5.5)

    # 4.7.2
    add_heading_styled(doc, "4.7.2 Use Case", 3)
    add_body(
        doc,
        "A Use Case Diagram is a UML diagram that illustrates the interactions between actors and "
        "a system, showing the services (use cases) the system provides and who uses them. Figure 7 "
        "below shows the use case diagram for the System for Personal Computer Checks Using QR "
        "Codes, depicting the interactions between gate officers, administrators, and the system.",
    )
    add_figure(doc, "Figure7_Proposed_UseCase.png", "Figure 7: Use case diagram of the proposed system", width=5.5)

    # 4.7.3 Normalization
    add_heading_styled(doc, "4.7.3 Normalization", 3)
    add_body(
        doc,
        "Normalization is the process of organizing data in a database to reduce data redundancy "
        "(duplication) and improve data integrity. It involves dividing large tables into smaller, "
        "related tables and defining relationships between them. Its objectives are to eliminate "
        "data redundancy, avoid data inconsistency, improve data integrity, simplify database "
        "maintenance, and reduce update, insertion, and deletion anomalies. The Computer_info, "
        "Users, and Logs tables of the proposed system were assessed against the following normal forms.",
    )

    add_heading_styled(doc, "4.7.3.1 First Normal Form (1NF)", 4)
    add_body(
        doc,
        "A table is in 1NF if each column contains atomic (indivisible) values, there are no "
        "repeating groups or multivalued attributes, and each record is unique. The Computer_info, "
        "Users, and Logs tables satisfy 1NF, as each field (for example sn, model, owno, owname, "
        "and date) holds a single, indivisible value, and each record is uniquely identified by a "
        "primary key (id or log_id).",
    )

    add_heading_styled(doc, "4.7.3.2 Second Normal Form (2NF)", 4)
    add_body(
        doc,
        "A table is in 2NF if it is already in 1NF and all non-key attributes are fully dependent "
        "on the entire primary key. Since each of the three tables uses a single-column primary "
        "key (id or log_id) rather than a composite key, every non-key attribute is fully dependent "
        "on that key, satisfying 2NF.",
    )

    add_heading_styled(doc, "4.7.3.3 Third Normal Form (3NF)", 4)
    add_body(
        doc,
        "A table is in 3NF if it is already in 2NF and there are no transitive dependencies, "
        "meaning non-key attributes should not depend on other non-key attributes. In the proposed "
        "design, attributes such as owname and owno depend directly on the device record identified "
        "by id, and not on any other non-key attribute, satisfying 3NF.",
    )

    add_heading_styled(doc, "4.7.3.4 Boyce-Codd Normal Form (BCNF)", 4)
    add_body(
        doc,
        "BCNF is a stronger version of 3NF: a table is in BCNF if, for every functional dependency, "
        "the determinant is a candidate key. BCNF eliminates certain anomalies that may still exist "
        "in 3NF. The Computer_info, Users, and Logs tables were reviewed and found to meet this "
        "condition, since sn and nid, which could otherwise determine other attributes, are "
        "themselves defined as unique keys.",
    )

    # 4.7.4 Data Dictionary
    add_heading_styled(doc, "4.7.4 Data Dictionary", 3)
    add_body(
        doc,
        "A data dictionary, also called a metadata repository, is a centralized description of the "
        "data used in a system. It explains what each piece of data means, how it is stored, how it "
        "is used, and how it relates to other data. A data dictionary is important because it helps "
        "developers, system administrators, and system users understand the database structure and "
        "maintain consistency across the system.",
    )
    add_body(
        doc,
        "Since the System for Personal Computer Checks Using QR Codes uses a database to store "
        "device, user, and activity-log information, the data dictionary below describes the main "
        "entities used in the system and their attributes. The entities include: Computer_info, "
        "Users, and Logs.",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Table 3: Computer_info table")
    set_run_font(r, size=11, bold=True)
    add_table(
        doc,
        ["Field", "Meaning", "Data Type", "Size", "Description / Function"],
        [
            ["id", "Device identification", "Int", "50", "Primary key, uniquely identifies each registered computer"],
            ["sn", "Serial number", "varchar", "20", "Unique identifier printed/encoded on the device's QR code"],
            ["model", "Device model", "varchar", "30", "Stores the model details of the computer"],
            ["owno", "Owner's number", "varchar", "30", "Stores the owner's contact number"],
            ["owname", "Owner's name", "varchar", "30", "Stores the name of the device owner"],
            ["date", "Date of record", "Date", "—", "Stores the date the device was registered"],
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Table 4: Users table")
    set_run_font(r, size=11, bold=True)
    add_table(
        doc,
        ["Field", "Meaning", "Data Type", "Size", "Description / Function"],
        [
            ["id", "User identification", "Int", "10", "Primary key, uniquely identifies each system user"],
            ["user_type", "Role", "varchar", "30", "Stores the user's role (Administrator, Gate Officer)"],
            ["nid", "User's number", "varchar", "20", "Unique national/staff identification number"],
            ["names", "Names", "varchar", "50", "Stores the full name of the user"],
            ["email", "Email", "varchar", "60", "Used for login and user identification"],
            ["password", "Password", "varchar", "200", "Stores the user's hashed password"],
        ],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Table 5: Logs table")
    set_run_font(r, size=11, bold=True)
    add_table(
        doc,
        ["Field", "Meaning", "Data Type", "Size", "Description / Function"],
        [
            ["log_id", "Log identification", "Int", "11", "Primary key, uniquely identifies each log entry"],
            ["sn", "Serial number", "varchar", "20", "References the device's serial number"],
            ["model", "Model", "varchar", "40", "Stores the model of the device involved in the action"],
            ["owno", "Owner number", "varchar", "20", "Stores the owner's contact number"],
            ["owname", "Owner names", "varchar", "80", "Stores the name of the device owner"],
            ["action", "Action done", "varchar", "50", "Stores the action performed (e.g., registered, verified, exited)"],
            ["comment", "Short comment", "varchar", "100", "Stores any remark made by the gate officer"],
            ["date", "Current time", "timestamp", "—", "Stores the date and time the action occurred"],
        ],
    )

    # 4.7.5 ERD
    add_heading_styled(doc, "4.7.5 Entity Relationship Diagram", 3)
    add_body(
        doc,
        "An Entity Relationship Diagram (ERD) is a graphical representation that shows how data is "
        "organized in a database and how different entities are related. It helps in designing the "
        "database structure by identifying the main entities, their attributes, and the relationships "
        "between them.",
    )
    add_figure(doc, "Figure8_ERD.png", "Figure 8: Entity Relationship Diagram", width=6.2)
    add_body(
        doc,
        "The diagram above represents the ER model of the System for Personal Computer Checks Using "
        "QR Codes. It describes the main entities involved in the system — Users, Computer_info, "
        "and Logs — along with their relationships, such as a user performing an action that is "
        "recorded in the logs, and a log entry referencing the device it relates to. The ERD "
        "provides a clear understanding of how device and user information is stored and linked "
        "within the database, supporting efficient data management and consistent implementation "
        "of backend operations.",
        first_line_indent=False,
    )

    # 4.8 Architecture
    add_heading_styled(doc, "4.8 Architecture design of the new system", 2)
    add_heading_styled(doc, "4.8.1 System Architecture", 3)
    add_body(
        doc,
        "The proposed System for Personal Computer Checks Using QR Codes for UTB Rubavu Campus "
        "adopts a Three-Tier Architecture consisting of the Presentation Layer, Application Layer, "
        "and Database Layer. This architecture promotes scalability, maintainability, security, "
        "and efficient data management.",
    )
    add_figure(doc, "Figure9_System_Architecture.png", "Figure 9: Three-Tier System Architecture", width=3.2)

    add_body(doc, "1. Presentation Layer (Client Layer)", first_line_indent=False)
    add_body(
        doc,
        "This is the user interface through which gate officers and administrators interact with "
        "the system. It provides web pages for registering devices, generating and printing QR "
        "codes, and reviewing logs and reports. Technologies used are HTML, Bootstrap, and "
        "JavaScript, supporting the following functions:",
    )
    for b in [
        "User registration and login",
        "Registering a personal computer at entry",
        "Generating and printing a device's QR code",
        "Scanning and verifying a QR code at exit",
        "Viewing logs and reports",
    ]:
        add_bullet(doc, b)

    add_body(doc, "2. Application Layer (Business Logic Layer)", first_line_indent=False)
    add_body(
        doc,
        "This layer processes requests from users and applies business rules before interacting "
        "with the database. The technology used is PHP with PDO, and its functions include:",
    )
    for b in [
        "User authentication and authorization",
        "Device registration and record management",
        "QR code generation and verification logic",
        "Logging of actions performed at the gate",
        "Report generation",
        "Data validation",
    ]:
        add_bullet(doc, b)

    add_body(doc, "3. Database Layer (Data Layer)", first_line_indent=False)
    add_body(
        doc,
        "This layer stores and manages all system data. The technology used is the MySQL Database "
        "Management System, hosted locally through XAMPP, which stores data such as user accounts, "
        "registered computers, and activity logs.",
    )

    doc.save(DOC_PATH)
    print("Saved:", DOC_PATH)


if __name__ == "__main__":
    build()
