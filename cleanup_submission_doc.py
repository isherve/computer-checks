"""
Cleanup submission doc: remove leftover manual TOC junk,
apply proper Figure/Table SEQ captions, refresh Word fields.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from docx.text.paragraph import Paragraph

DOC_PATH = Path(r"C:\Users\ishim\OneDrive\Desktop\EMMY - QR\EMMY_QR_FINAL_SUBMISSION.docx")


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def delete_paragraph(paragraph: Paragraph):
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def clear_paragraph(paragraph: Paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_seq_field(paragraph: Paragraph, seq_name: str):
    """Insert SEQ field: { SEQ Figure \\* ARABIC }"""
    run = paragraph.add_run()
    r = run._r

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r.append(begin)

    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC "
    instr_run.append(instr)
    paragraph._p.append(instr_run)

    sep_run = paragraph.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(sep)

    # temporary result
    tmp = paragraph.add_run("1")
    set_run_font(tmp, size=11, bold=True, italic=True)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def make_caption(paragraph: Paragraph, kind: str, number: str, title: str):
    """Rebuild caption as: Figure {SEQ} : title  (Caption style)."""
    clear_paragraph(paragraph)
    try:
        paragraph.style = "Caption"
    except Exception:
        pass
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)

    label = "Figure" if kind.lower().startswith("fig") else "Table"
    r1 = paragraph.add_run(f"{label} ")
    set_run_font(r1, size=11, bold=True, italic=True)
    add_seq_field(paragraph, label)
    r2 = paragraph.add_run(f": {title}")
    set_run_font(r2, size=11, bold=True, italic=True)


def is_manual_toc_line(text: str) -> bool:
    t = text.strip()
    if not t:
        return False
    # tab-separated page entries
    if "\t" in t and re.search(r"\t+\.?\.?\s*\d+[ivx]*\s*$", t, re.I):
        return True
    if "\t" in t and re.search(r"\d+\s*$", t):
        # e.g. CHAPTER ONE...\t1
        return True
    # dotted leaders style remnants
    if re.search(r"\.{3,}\s*\d+\s*$", t):
        return True
    return False


def cleanup():
    doc = Document(str(DOC_PATH))

    # 1) Delete leftover manual TOC / LOF / LOT lines (not the Heading 1 titles themselves)
    # Also delete duplicate early CHAPTER headings that are only TOC shells
    # Strategy: after TABLE OF CONTENTS field block, before ABSTRACT content,
    # remove paragraphs that look like TOC entries OR duplicate chapter titles that appear
    # BEFORE the real ABSTRACT heading content body.

    # Find indices
    paras = list(doc.paragraphs)
    abstract_idx = None
    toc_idx = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == "TABLE OF CONTENTS" and toc_idx is None:
            toc_idx = i
        if t == "ABSTRACT" and abstract_idx is None:
            # prefer the one that is followed by abstract body, not toc
            abstract_idx = i

    to_delete = []
    seen_real_chapters = set()

    for i, p in enumerate(paras):
        t = p.text.strip()
        st = p.style.name if p.style else ""

        # Remove update-field notes
        if t.startswith("(Update this field"):
            to_delete.append(p)
            continue

        # Remove manual TOC-like lines everywhere in front matter / old TOC
        if is_manual_toc_line(t):
            to_delete.append(p)
            continue

        # Remove orphan 'LITERATURE REVIEW' heading if CHAPTER TWO already includes it
        if t.upper() == "LITERATURE REVIEW" and st.startswith("Heading"):
            to_delete.append(p)
            continue

        # Remove "No table of figures entries found." placeholders (will regenerate)
        if t.lower().startswith("no table of figures entries found"):
            # keep field paras - actually leave for refresh; Word put this as field result
            continue

    for p in to_delete:
        try:
            delete_paragraph(p)
        except Exception:
            pass

    # 2) Deduplicate Heading 1 chapter titles: keep last occurrence of each unique chapter title
    # (real content is later in the document). Delete earlier duplicates.
    paras = list(doc.paragraphs)
    chapter_positions = {}
    for i, p in enumerate(paras):
        t = p.text.strip().upper()
        st = p.style.name if p.style else ""
        if st == "Heading 1" and t.startswith("CHAPTER"):
            chapter_positions.setdefault(t, []).append(p)

    for title, plist in chapter_positions.items():
        if len(plist) > 1:
            # keep the last one (actual chapter body), delete earlier
            for p in plist[:-1]:
                delete_paragraph(p)

    # Also dedupe ABSTRACT / REFERENCES / APPENDICES Heading 1 if duplicated
    for key in ("ABSTRACT", "REFERENCES", "APPENDICES", "ACKNOWLEDGEMENT",
                "LIST OF FIGURES", "LIST OF TABLES", "TABLE OF CONTENTS",
                "LIST OF SYMBOLS, ACRONYMS AND ABBREVIATIONS"):
        hits = [p for p in doc.paragraphs if p.text.strip().upper() == key and (p.style and p.style.name == "Heading 1")]
        if len(hits) > 1:
            # For front-matter titles keep first; for REFERENCES/APPENDICES/ABSTRACT keep appropriate
            if key in ("REFERENCES", "APPENDICES"):
                for p in hits[:-1]:
                    delete_paragraph(p)
            elif key == "ABSTRACT":
                # keep the one near body - usually last before ch1 body; keep last
                for p in hits[:-1]:
                    delete_paragraph(p)
            else:
                for p in hits[1:]:
                    delete_paragraph(p)

    # 3) Convert captions to SEQ-based captions
    caption_re = re.compile(r"^(Figure|Table)\s+(\d+)\s*:\s*(.+)$", re.I)
    for p in doc.paragraphs:
        t = p.text.strip()
        m = caption_re.match(t)
        if not m:
            continue
        # skip narrative
        if "presents a" in t.lower() or t.lower().startswith("additionally"):
            continue
        kind, num, title = m.group(1), m.group(2), m.group(3).strip()
        # avoid very long narrative captions
        if len(title) > 120 and ("flowchart" in title.lower() and "presents" in title.lower()):
            continue
        make_caption(p, kind, num, title)

    # 4) Front-matter titles should NOT use Heading 1 (avoids TOC pollution).
    front_titles = {
        "LIST OF FIGURES": "LIST OF FIGURES",
        "LIST OF TABLES": "LIST OF TABLES",
        "TABLE OF CONTENTS": "TABLE OF CONTENTS",
        "ACKNOWLEDGEMENT": "ACKNOWLEDGEMENT",
        "LIST OF SYMBOLS, ACRONYMS AND ABBREVIATIONS": "LIST OF SYMBOLS, ACRONYMS AND ABBREVIATIONS",
        "DECLARATION": "DECLARATION",
        "CERTIFICATE": "CERTIFICATE",
        "DEDICATION": "DEDICATION",
    }
    for p in doc.paragraphs:
        t = p.text.strip().upper()
        if t in front_titles:
            clear_paragraph(p)
            try:
                p.style = "Normal"
            except Exception:
                pass
            run = p.add_run(front_titles[t])
            set_run_font(run, size=14, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._p.get_or_add_pPr()
            if not any(c.tag == qn("w:pageBreakBefore") for c in pPr):
                pPr.append(OxmlElement("w:pageBreakBefore"))
            continue

        if p.style and p.style.name == "Heading 1":
            if t.startswith("CHAPTER") or t in ("ABSTRACT", "REFERENCES", "APPENDICES"):
                pPr = p._p.get_or_add_pPr()
                if not any(c.tag == qn("w:pageBreakBefore") for c in pPr):
                    pPr.append(OxmlElement("w:pageBreakBefore"))
                for run in p.runs:
                    set_run_font(run, size=14, bold=True)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5) Fix hanging refs formatting
    in_refs = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.upper() == "REFERENCES":
            in_refs = True
            continue
        if t.upper().startswith("APPENDIX"):
            in_refs = False
        if in_refs and t.startswith("["):
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.first_line_indent = Cm(-0.75)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                set_run_font(run, size=11)

    doc.save(str(DOC_PATH))
    print("Cleanup saved.")


def update_word_fields():
    import win32com.client

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(DOC_PATH))
        # Update SEQ fields first, then TOCs
        doc.Fields.Update()
        for toc in doc.TablesOfContents:
            toc.Update()
        # TablesOfFigures may be empty collection type - try
        try:
            for i in range(1, doc.TablesOfFigures.Count + 1):
                doc.TablesOfFigures(i).Update()
        except Exception:
            pass
        # Repaginate
        doc.Repaginate()
        doc.Fields.Update()
        doc.Save()
        doc.Close(False)
        print("Word fields refreshed.")
    finally:
        word.Quit()


if __name__ == "__main__":
    cleanup()
    update_word_fields()
