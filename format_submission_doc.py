"""
Format emmy new.docx for submission:
- Clean structure / headings / captions
- IEEE in-text citations + References
- Automatic TOC, List of Figures, List of Tables (Word fields)
"""
from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.text.paragraph import Paragraph

SRC = Path(r"C:\Users\ishim\OneDrive\Desktop\EMMY - QR\emmy new.docx")
DST = Path(r"C:\Users\ishim\OneDrive\Desktop\EMMY - QR\EMMY_QR_FINAL_SUBMISSION.docx")

IEEE_REFS = [
    '[1] Wikipedia, "Data Dictionary," 2024. [Online]. Available: https://en.wikipedia.org/wiki/Data_dictionary',
    '[2] F. Almasalha, "QR code based applications," Jan. 2014. [Online].',
    '[3] M.-J. Tsai, "QR code beautification by instance segmentation (IS-QR)," Digital Signal Processing, 2024.',
    '[4] S. Tiwari, "An Introduction To QR Code Technology," in Proc. 2016 Int. Conf. Information Technology (InCITe), 2016.',
    '[5] S. Pati, A Novel QR Code Based Smart Attendance Tracking System. Kolkata, India: Dept. of ECE, 2020.',
    '[6] Malwarebytes, "QR Codes: How they work and how to stay safe," 2024. [Online]. Available: https://www.malwarebytes.com/',
    '[7] M. Innocent and U. Jeannine, Gate Check Management System. Kigali, Rwanda, 2020, p. 11.',
    '[8] GeeksforGeeks, "What is DFD (Data Flow Diagram)?," 2024. [Online]. Available: https://www.geeksforgeeks.org/what-is-dfd/',
    '[9] R. Elmasri and S. B. Navathe, Fundamentals of Database Systems, 7th ed. Pearson, 2016.',
    '[10] I. Sommerville, Software Engineering, 10th ed. Pearson, 2016.',
    '[11] T. Connolly and C. Begg, Database Systems: A Practical Approach to Design, Implementation, and Management, 6th ed. Pearson, 2015.',
    '[12] PHP Group, "PHP Data Objects (PDO)," PHP Manual, 2024. [Online]. Available: https://www.php.net/manual/en/book.pdo.php',
]


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def delete_paragraph(paragraph: Paragraph):
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_paragraph_after(paragraph: Paragraph, text="", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def add_page_break_before(paragraph: Paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    # remove existing pageBreakBefore if any then set
    for child in list(pPr):
        if child.tag == qn("w:pageBreakBefore"):
            pPr.remove(child)
    pb = OxmlElement("w:pageBreakBefore")
    pPr.append(pb)


def clear_paragraph(paragraph: Paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, *, size=12, bold=False, italic=False, align=None, style=None):
    clear_paragraph(paragraph)
    if style:
        try:
            paragraph.style = style
        except Exception:
            pass
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return run


def add_field(paragraph: Paragraph, instr: str):
    """Insert a Word complex field into an existing paragraph."""
    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_run = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    instr_run.append(instr_text)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    # placeholder text shown before update
    placeholder = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click → Update Field (or press Ctrl+A then F9 in Word)"
    placeholder.append(t)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    paragraph._p.append(instr_run)
    run2 = paragraph.add_run()
    run2._r.append(fld_sep)
    paragraph._p.append(placeholder)
    run3 = paragraph.add_run()
    run3._r.append(fld_end)


def ensure_caption_style(doc: Document):
    try:
        style = doc.styles["Caption"]
    except KeyError:
        style = doc.styles.add_style("Caption", 1)  # paragraph style
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.font.italic = True
    style.font.bold = True
    return style


def format_body_paragraph(paragraph: Paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    for run in paragraph.runs:
        if not run.font.name:
            set_run_font(run)
        else:
            # force TNR for submission consistency
            set_run_font(run, size=run.font.size.pt if run.font.size else 12,
                         bold=bool(run.bold), italic=bool(run.italic))


CHAPTER_HEADING_MAP = {
    r"^CHAPTER\s+ONE[:\s].*": ("CHAPTER ONE: GENERAL INTRODUCTION", 1),
    r"^CHAPTER\s+TWO\b.*": ("CHAPTER TWO: LITERATURE REVIEW", 1),
    r"^LITERATURE REVIEW$": None,  # merge handled separately
    r"^CHAPTER\s+THREE.*": ("CHAPTER THREE: RESEARCH METHODOLOGY", 1),
    r"^CHAPTER\s+FOUR.*": ("CHAPTER FOUR: SYSTEM ANALYSIS AND DESIGN", 1),
    r"^CHAPTER\s*5[:\s].*|^CHAPTER\s+FIVE.*": ("CHAPTER FIVE: IMPLEMENTATION, CODING AND TESTING", 1),
    r"^CHAPTER\s+SIX.*": ("CHAPTER SIX: CONCLUSION AND RECOMMENDATIONS", 1),
    r"^REFERENCES$": ("REFERENCES", 1),
    r"^APPENDIXES$|^APPENDICES$": ("APPENDICES", 1),
    r"^ABSTRACT$": ("ABSTRACT", 1),
    r"^ACKNOWLEDGEMENT$|^ACKNOWLEDGMENT$": ("ACKNOWLEDGEMENT", 1),
    r"^LIST OF SYMBOLS.*": ("LIST OF SYMBOLS, ACRONYMS AND ABBREVIATIONS", 1),
    r"^LIST OF FIGURES$": ("LIST OF FIGURES", 1),
    r"^LIST OF TABLES$": ("LIST OF TABLES", 1),
    r"^TABLE OF CONTENTS$": ("TABLE OF CONTENTS", 1),
}


SECTION_REWRITES = {
    "INTRODUCTION": None,  # contextual
    "1,6 SCOPE OF THE PROJECT": "1.6 Scope of the Project",
    "1.7 ORGANIZATION OF THE PROJECT": "1.7 Organization of the Project",
    "1.8 ORGANIZATION OF THE PROJECT": "1.7 Organization of the Project",
    "PROBLEM STATEMENT": "1.3 Problem Statement",
    "OBJECTIVES": "1.4 Objectives",
    "General objective": "1.4.1 General Objective",
    "Specific objectives": "1.4.2 Specific Objectives",
    "Research questions": "1.4.3 Research Questions",
    "Background to the study": "1.2 Background of the Study",
    "CONCLUSION": "6.1 Conclusion",
    "RECOMMENDATIONS": "6.2 Recommendations",
    "Key Terms": "Key Terms",
}


def classify_heading(text: str):
    t = text.strip()
    if re.match(r"^Figure\s+\d+\s*:", t, re.I):
        return "caption_figure"
    if re.match(r"^Table\s+\d+\s*:", t, re.I):
        return "caption_table"
    # numbered sections
    if re.match(r"^\d+\.\d+(\.\d+)*(\.\d+)*\s+\S", t):
        # count dots for level
        num = t.split()[0]
        depth = num.count(".") + 1
        level = min(depth + 1, 4)  # Heading 2..4 roughly; keep Heading2 for x.y
        if depth == 1:
            return ("heading", 2, t)
        if depth == 2:
            return ("heading", 3, t)
        return ("heading", 4, t)
    if re.match(r"^2\.\d+", t) or re.match(r"^3\.\d+", t) or re.match(r"^4\.\d+", t) or re.match(r"^5\.\d+", t) or re.match(r"^6\.\d+", t) or re.match(r"^1\.\d+", t):
        return ("heading", 2, t)
    return None


def apply_ieee_citations(text: str) -> str:
    """Add IEEE citation markers where literature claims appear (idempotent-ish)."""
    replacements = [
        (r"(QR code \(abbreviated from Quick Response Code\)[^.]*Japan\.)", r"\1 [4]"),
        (r"(QR code technology has become one of the most widely adopted identification technologies[^.]*\.)", r"\1 [3], [4]"),
        (r"(Computer asset management systems are designed to monitor organizational computer resources[^.]*\.)", r"\1 [5], [7]"),
        (r"(Modern asset management systems depend on centralized databases[^.]*\.)", r"\1 [9], [11]"),
        (r"(Reporting systems transform stored operational data into meaningful reports[^.]*\.)", r"\1 [10]"),
        (r"(Many educational institutions continue to rely on manual recording methods[^.]*\.)", r"\1 [7]"),
        (r"(A Quick Response \(QR\) code is a two-dimensional barcode[^.]*\.)", r"\1 [4], [6]"),
        (r"(A Database Management System \(DBMS\) is software used to create, manage[^.]*\.)", r"\1 [9]"),
        (r"(A Data Flow Diagram \(DFD\) is a graphical representation of how data moves through a system\.[^\[]*)",
         r"\1 [8]"),
        (r"(The conceptual model demonstrates how computer information flows[^.]*\.)", r"\1 [10]"),
        (r"(Normalization is the process of organizing data in a database[^.]*\.)", r"\1 [9], [11]"),
        (r"(A data dictionary, also called a metadata repository[^.]*\.)", r"\1 [1]"),
        (r"(PHP with PDO[^.]*prepared statements[^.]*\.)", r"\1 [12]"),
    ]
    out = text
    for pat, repl in replacements:
        if re.search(r"\[\d+\]", out) and re.search(pat, out):
            # already has a citation nearby in same sentence - skip if already cited for this pattern
            m = re.search(pat, out)
            if m and re.search(r"\[\d+\]", m.group(0)):
                continue
        # avoid double-adding identical cites
        new = re.sub(pat, repl, out, count=1)
        # prevent duplicated brackets like [4] [4]
        new = re.sub(r"(\[\d+\])(\s+\1)+", r"\1", new)
        out = new
    return out


def rebuild_front_matter_fields(doc: Document):
    """Replace manual LOF/LOT/TOC entries with automatic field placeholders."""
    # Find the heading paragraphs for LOF, LOT, TOC (real body ones, not TOC lines with tabs)
    targets = {}
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "LIST OF FIGURES" and "\t" not in p.text:
            targets["lof"] = (i, p)
        elif t == "LIST OF TABLES" and "\t" not in p.text:
            targets["lot"] = (i, p)
        elif t == "TABLE OF CONTENTS" and "\t" not in p.text:
            targets["toc"] = (i, p)

    def wipe_following_manual_entries(start_para, stop_titles):
        # delete subsequent paragraphs until we hit another major title
        nxt = start_para._element.getnext()
        removed = 0
        while nxt is not None:
            # only care about paragraphs
            if nxt.tag != qn("w:p"):
                break
            txt = "".join(node.text or "" for node in nxt.iter(qn("w:t"))).strip()
            if any(txt.upper().startswith(s) for s in stop_titles):
                break
            # keep blank? remove manual numbered entries / list paragraphs
            if txt == "" and removed > 8:
                break
            following = nxt.getnext()
            parent = nxt.getparent()
            parent.remove(nxt)
            nxt = following
            removed += 1
            if removed > 80:
                break

    if "lof" in targets:
        _, p = targets["lof"]
        set_paragraph_text(p, "LIST OF FIGURES", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, style="Heading 1")
        add_page_break_before(p)
        wipe_following_manual_entries(p, ["LIST OF TABLES", "TABLE OF CONTENTS", "DECLARATION", "ABSTRACT", "CHAPTER"])
        field_p = insert_paragraph_after(p, "")
        add_field(field_p, r'TOC \h \z \c "Figure"')
        note = insert_paragraph_after(field_p, "")
        set_paragraph_text(note, "(Update this field in Microsoft Word: select all → press F9)", size=10, italic=True)

    if "lot" in targets:
        # re-find after mutations
        for p in doc.paragraphs:
            if p.text.strip() == "LIST OF TABLES":
                set_paragraph_text(p, "LIST OF TABLES", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, style="Heading 1")
                add_page_break_before(p)
                wipe_following_manual_entries(p, ["TABLE OF CONTENTS", "DECLARATION", "ABSTRACT", "CHAPTER", "LIST OF"])
                field_p = insert_paragraph_after(p, "")
                add_field(field_p, r'TOC \h \z \c "Table"')
                note = insert_paragraph_after(field_p, "")
                set_paragraph_text(note, "(Update this field in Microsoft Word: select all → press F9)", size=10, italic=True)
                break

    # TOC - find heading then wipe old manual TOC
    for p in doc.paragraphs:
        if p.text.strip() == "TABLE OF CONTENTS":
            set_paragraph_text(p, "TABLE OF CONTENTS", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, style="Heading 1")
            add_page_break_before(p)
            wipe_following_manual_entries(p, ["ABSTRACT", "CHAPTER ONE", "CHAPTER 1", "DECLARATION", "KEY TERMS"])
            # Also remove leftover DECLARATION...REFERENCES manual TOC lines that start right after
            # (already wiped)
            field_p = insert_paragraph_after(p, "")
            add_field(field_p, r'TOC \o "1-3" \h \z \u')
            note = insert_paragraph_after(field_p, "")
            set_paragraph_text(note, "(Update this field in Microsoft Word: select all → press F9)", size=10, italic=True)
            break


def convert_captions(doc: Document):
    ensure_caption_style(doc)
    for p in doc.paragraphs:
        t = p.text.strip()
        if re.match(r"^Figure\s+\d+\s*:", t, re.I) or re.match(r"^Table\s+\d+\s*:", t, re.I):
            # Don't convert narrative sentences like "Figure 2 presents..."
            if len(t) > 120 and not re.match(r"^(Figure|Table)\s+\d+\s*:\s*.{1,80}$", t):
                # if it's "Figure X: Short caption" ok; long text may still be caption
                if "presents a" in t.lower() or "below" in t.lower() and ":" not in t[:20]:
                    continue
            if re.match(r"^Figure\s+\d+\s*presents", t, re.I):
                continue
            if re.match(r"^Additionally,\s*Figure", t, re.I):
                continue
            try:
                p.style = "Caption"
            except Exception:
                pass
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            for run in p.runs:
                set_run_font(run, size=11, bold=True, italic=True)


def normalize_headings(doc: Document):
    merge_chapter_two = False
    for i, p in enumerate(doc.paragraphs):
        raw = p.text
        t = raw.strip()
        if not t:
            continue

        # Remove AI meta commentary
        if t.startswith("Yes. The second format is far more academic") or t.startswith('For your project "Design and Implementation'):
            delete_paragraph(p)
            continue

        # Fix junk heading
        if t == "ed" and p.style and "Heading" in p.style.name:
            delete_paragraph(p)
            continue

        # Fix typos in place
        fixed = t
        fixed = fixed.replace("GENRAL", "GENERAL")
        fixed = fixed.replace("CERTFICATE", "CERTIFICATE")
        fixed = re.sub(r"^1,6 SCOPE", "1.6 Scope", fixed)
        if fixed != t:
            set_paragraph_text(p, fixed, size=12, bold=("Heading" in (p.style.name if p.style else "")))
            t = fixed

        # Chapter TWO split over two Heading 1 lines
        if t.upper() == "CHAPTER TWO":
            # peek next
            nxt = None
            for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip():
                    nxt = doc.paragraphs[j]
                    break
            if nxt and nxt.text.strip().upper() == "LITERATURE REVIEW":
                set_paragraph_text(p, "CHAPTER TWO: LITERATURE REVIEW", size=14, bold=True,
                                   align=WD_ALIGN_PARAGRAPH.CENTER, style="Heading 1")
                add_page_break_before(p)
                delete_paragraph(nxt)
                continue

        # Major chapter titles
        for pat, val in CHAPTER_HEADING_MAP.items():
            if re.match(pat, t, re.I):
                if val is None:
                    continue
                title, level = val
                set_paragraph_text(p, title, size=14 if level == 1 else 12, bold=True,
                                   align=WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT,
                                   style=f"Heading {level}")
                if level == 1 and title.startswith("CHAPTER"):
                    add_page_break_before(p)
                break

        # Section rewrites for Ch1 / Ch6 list-style headings
        if t in SECTION_REWRITES and SECTION_REWRITES[t]:
            new_t = SECTION_REWRITES[t]
            level = 3 if re.match(r"^\d+\.\d+\.\d+", new_t) else 2
            set_paragraph_text(p, new_t, size=12, bold=True, style=f"Heading {level}")
            continue

        if t.upper() == "INTRODUCTION" and i < 140:
            set_paragraph_text(p, "1.1 Introduction", size=12, bold=True, style="Heading 2")
            continue

        # Already numbered academic headings
        cls = classify_heading(t)
        if cls == "caption_figure" or cls == "caption_table":
            continue
        if isinstance(cls, tuple) and cls[0] == "heading":
            _, level, text = cls
            # Keep existing good headings but ensure style
            try:
                p.style = f"Heading {level}"
            except Exception:
                pass
            for run in p.runs:
                set_run_font(run, size=12, bold=True)


def add_citations_to_body(doc: Document):
    for p in doc.paragraphs:
        t = p.text
        if not t.strip():
            continue
        # skip headings/captions/code-ish
        style = p.style.name if p.style else ""
        if "Heading" in style or style == "Caption":
            continue
        if t.strip().startswith("<") or t.strip().startswith("$") or t.strip().startswith("//"):
            continue
        new = apply_ieee_citations(t)
        if new != t:
            # preserve roughly by rewriting
            align = p.alignment
            set_paragraph_text(p, new, size=12)
            if align:
                p.alignment = align
            format_body_paragraph(p)


def rebuild_references(doc: Document):
    # Find REFERENCES heading
    ref_p = None
    for p in doc.paragraphs:
        if p.text.strip().upper() in ("REFERENCES", "REFERENCE"):
            ref_p = p
            break
    if not ref_p:
        return

    set_paragraph_text(ref_p, "REFERENCES", size=14, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, style="Heading 1")
    add_page_break_before(ref_p)

    # Remove old reference table if present near end
    # (python-docx can't easily delete tables by proximity; clear following paras until APPENDIX)
    nxt = ref_p._element.getnext()
    while nxt is not None:
        if nxt.tag == qn("w:tbl"):
            following = nxt.getnext()
            nxt.getparent().remove(nxt)
            nxt = following
            continue
        if nxt.tag == qn("w:p"):
            txt = "".join(node.text or "" for node in nxt.iter(qn("w:t"))).strip().upper()
            if txt.startswith("APPENDIX"):
                break
            if txt == "":
                following = nxt.getnext()
                nxt.getparent().remove(nxt)
                nxt = following
                continue
            # remove old ref lines that look like [n]
            if re.match(r"^\[\d+\]", txt) or "wikipedia" in txt.lower() or "available:" in txt.lower():
                following = nxt.getnext()
                nxt.getparent().remove(nxt)
                nxt = following
                continue
            # if we hit unexpected content that's appendix code, stop only on APPENDIX
            if txt.startswith("<!DOCTYPE") or txt.startswith("BELOW ARE SOME CODES"):
                break
        break

    # Insert IEEE references after heading
    anchor = ref_p
    for ref in IEEE_REFS:
        anchor = insert_paragraph_after(anchor, "")
        set_paragraph_text(anchor, ref, size=11)
        anchor.paragraph_format.left_indent = Cm(0.75)
        anchor.paragraph_format.first_line_indent = Cm(-0.75)  # hanging indent
        anchor.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        anchor.paragraph_format.space_after = Pt(6)
        anchor.paragraph_format.line_spacing = 1.15


def format_all_body(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)

    for p in doc.paragraphs:
        style = p.style.name if p.style else "Normal"
        t = p.text.strip()
        if not t:
            continue
        if "Heading" in style:
            for run in p.runs:
                set_run_font(run, size=14 if style == "Heading 1" else 12, bold=True)
            if style == "Heading 1":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
            continue
        if style == "Caption":
            continue
        # skip pure code lines
        if t.startswith("<") or t.startswith("$") or t.startswith("//") or t.startswith("QR/"):
            for run in p.runs:
                set_run_font(run, name="Consolas", size=9)
            continue
        if style in ("Normal", "Body Text", "Normal (Web)", "List Paragraph", "List Bullet"):
            format_body_paragraph(p)


def update_fields_with_word(path: Path) -> bool:
    try:
        import win32com.client  # type: ignore
    except Exception:
        return False
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(path))
        # Update all fields including TOC
        doc.Fields.Update()
        # Also update TOCs collection
        for toc in doc.TablesOfContents:
            toc.Update()
        for tof in doc.TablesOfFigures:
            tof.Update()
        doc.Save()
        doc.Close(False)
        return True
    except Exception as e:
        print("Word update failed:", e)
        return False
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def main():
    print("Copying source...")
    shutil.copy2(SRC, DST)
    doc = Document(str(DST))

    print("Normalizing headings / cleaning...")
    normalize_headings(doc)

    print("Converting captions...")
    convert_captions(doc)

    print("Adding IEEE citations...")
    add_citations_to_body(doc)

    print("Rebuilding REFERENCES...")
    rebuild_references(doc)

    print("Inserting automatic TOC / LOF / LOT fields...")
    rebuild_front_matter_fields(doc)

    print("Formatting body...")
    format_all_body(doc)

    doc.save(str(DST))
    print("Saved:", DST)

    print("Updating fields via Word COM (if available)...")
    if update_fields_with_word(DST):
        print("Fields updated successfully.")
    else:
        print("Open the document in Word and press Ctrl+A then F9 to refresh TOC/LOF/LOT.")


if __name__ == "__main__":
    main()
