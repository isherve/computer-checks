"""
Configure page numbers + Heading 1/2/3 for TOC:
- Roman numerals (i, ii, iii...) before Chapter One
- Arabic numerals (1, 2, 3...) from Chapter One to end
- Normalize headings so subtitles appear in TOC (levels 1-3)
"""
from __future__ import annotations

import re
from pathlib import Path

import win32com.client
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

SRC = Path(r"C:\Users\ishim\OneDrive\Desktop\EMMY - QR\EMMY_QR_FINAL_SUBMISSION.docx")
# Also refresh the open working copy name if needed
OUT = SRC


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def normalize_headings_python():
    """Fix heading levels with python-docx before Word page-number setup."""
    doc = Document(str(SRC))

    def heading_level_for(text: str):
        t = text.strip()
        up = t.upper()
        if up.startswith("CHAPTER ") or up in ("ABSTRACT", "REFERENCES", "APPENDICES", "APPENDIXES"):
            return 1
        # 1.2.3 ... -> H3
        m = re.match(r"^(\d+(?:\.\d+)+)\b", t)
        if not m:
            return None
        depth = m.group(1).count(".")
        if depth == 1:  # 1.1, 2.0, 4.5
            return 2
        # 1.4.1, 4.6.1, 2.1.1 -> H3 (cap at 3 for TOC)
        return 3

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        # skip TOC / LOF generated entries
        st = p.style.name if p.style else ""
        if st.startswith("toc") or st.startswith("table of figures"):
            continue
        if t in ("LIST OF FIGURES", "LIST OF TABLES", "TABLE OF CONTENTS",
                 "ACKNOWLEDGEMENT", "LIST OF SYMBOLS, ACRONYMS AND ABBREVIATIONS",
                 "DECLARATION", "CERTIFICATE", "DEDICATION"):
            continue

        lvl = heading_level_for(t)
        if lvl is None:
            continue
        try:
            p.style = f"Heading {lvl}"
        except Exception:
            continue
        # Title-case fix for scope
        if re.match(r"^1\.6\s+Scope", t, re.I):
            # rewrite text cleanly
            for child in list(p._p):
                if child.tag != qn("w:pPr"):
                    p._p.remove(child)
            run = p.add_run("1.6 Scope of the Project")
            set_run_font(run, size=12, bold=True)

        for run in p.runs:
            set_run_font(run, size=14 if lvl == 1 else 12, bold=True)
        if lvl == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(OUT))
    print("Headings normalized (H1/H2/H3).")


def setup_page_numbers_and_toc():
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    # Word constants
    wdSectionBreakNextPage = 2
    wdAlignPageNumberCenter = 1
    wdOrientPortrait = 0
    wdHeaderFooterPrimary = 1
    wdSeekMainDocument = 0
    wdFieldEmpty = -1
    wdDoNotSaveChanges = 0

    try:
        doc = word.Documents.Open(str(OUT))

        # --- Find CHAPTER ONE body paragraph ---
        ch1 = None
        for i in range(1, doc.Paragraphs.Count + 1):
            p = doc.Paragraphs(i)
            txt = p.Range.Text.replace("\r", "").replace("\x07", "").strip()
            style = ""
            try:
                style = p.Style.NameLocal
            except Exception:
                pass
            if txt.upper().startswith("CHAPTER ONE") and "Heading" in str(style):
                ch1 = p
                break
            if txt.upper() == "CHAPTER ONE: GENERAL INTRODUCTION":
                ch1 = p
                # prefer heading but accept first real body occurrence after abstract
                # continue searching for Heading version
                if "Heading" in str(style):
                    break

        if ch1 is None:
            raise RuntimeError("CHAPTER ONE heading not found")

        print("Found CHAPTER ONE at style:", ch1.Style.NameLocal)

        # --- Ensure a section break exists immediately before Chapter One ---
        # If Chapter One is not at the start of a section, insert Next Page section break
        ch1_start = ch1.Range.Start
        # Move to start of paragraph and insert break before it
        # Check if already at section start (approximately)
        sec_before = doc.Range(0, ch1_start).Sections.Count

        # Insert section break before Chapter One (safe even if extra sections exist)
        rng = ch1.Range.Duplicate
        rng.Collapse(1)  # wdCollapseStart = 1
        # Only insert if previous char is not already a section break nuance:
        # Always insert Next Page section break before Chapter One for clean restart
        # First check: is Chapter One the first para of its section?
        sec = ch1.Range.Sections(1)
        first_para_start = sec.Range.Paragraphs(1).Range.Start
        if abs(first_para_start - ch1.Range.Start) > 5:
            rng.InsertBreak(wdSectionBreakNextPage)
            print("Inserted Next Page section break before CHAPTER ONE")
            # Re-find Chapter One after insert
            ch1 = None
            for i in range(1, doc.Paragraphs.Count + 1):
                p = doc.Paragraphs(i)
                txt = p.Range.Text.replace("\r", "").strip()
                try:
                    style = p.Style.NameLocal
                except Exception:
                    style = ""
                if txt.upper().startswith("CHAPTER ONE") and "Heading" in str(style):
                    ch1 = p
                    break
            if ch1 is None:
                raise RuntimeError("CHAPTER ONE lost after section break")

        body_section_index = ch1.Range.Sections(1).Index  # 1-based
        print(f"Body starts at section {body_section_index} of {doc.Sections.Count}")

        # --- Configure page numbers for each section ---
        for i in range(1, doc.Sections.Count + 1):
            section = doc.Sections(i)
            footer = section.Footers(wdHeaderFooterPrimary)
            footer.LinkToPrevious = False

            # Clear existing footer fields/text
            footer.Range.Delete()

            # Add centered PAGE field
            footer.PageNumbers.Add(
                PageNumberAlignment=wdAlignPageNumberCenter,
                FirstPage=True,
            )

            if i < body_section_index:
                # Front matter: Roman numerals
                section.PageSetup.StartingNumber = 1 if i == 1 else section.PageSetup.StartingNumber
                # Restart numbering at first front section; continuous roman for rest of front
                if i == 1:
                    footer.PageNumbers.RestartNumberingAtSection = True
                    footer.PageNumbers.StartingNumber = 1
                else:
                    # continue roman sequence
                    footer.PageNumbers.RestartNumberingAtSection = False
                footer.PageNumbers.NumberStyle = 2  # wdPageNumberStyleLowercaseRoman = 2? 
                # Constants: wdPageNumberStyleArabic=0, UpperRoman=1, LowerRoman=2
                try:
                    section.Footers(wdHeaderFooterPrimary).PageNumbers.NumberStyle = 2
                except Exception:
                    pass
                # Also set via PageSetup
                section.PageSetup.PageNumberStyle = 2  # lower roman
            else:
                # Body: Arabic from 1 at Chapter One section, continuous after
                section.PageSetup.PageNumberStyle = 0  # Arabic
                try:
                    section.Footers(wdHeaderFooterPrimary).PageNumbers.NumberStyle = 0
                except Exception:
                    pass
                if i == body_section_index:
                    footer.PageNumbers.RestartNumberingAtSection = True
                    footer.PageNumbers.StartingNumber = 1
                    section.PageSetup.StartingNumber = 1
                else:
                    footer.PageNumbers.RestartNumberingAtSection = False

        # Apply NumberStyle again more carefully using PageNumbers collection
        for i in range(1, doc.Sections.Count + 1):
            section = doc.Sections(i)
            footer = section.Footers(wdHeaderFooterPrimary)
            footer.LinkToPrevious = False
            pn = footer.PageNumbers
            if i < body_section_index:
                pn.NumberStyle = 2  # lowercase roman
                if i == 1:
                    pn.RestartNumberingAtSection = True
                    pn.StartingNumber = 1
                else:
                    pn.RestartNumberingAtSection = False
            else:
                pn.NumberStyle = 0  # Arabic
                if i == body_section_index:
                    pn.RestartNumberingAtSection = True
                    pn.StartingNumber = 1
                else:
                    pn.RestartNumberingAtSection = False

        print("Page numbering configured.")

        # --- Unlink auto list numbering on headings (avoid 1.1 1.1 in TOC) ---
        for name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
            try:
                doc.Styles(name).LinkToListTemplate(Link=False)
            except Exception:
                pass

        # --- Rebuild TOC including Heading 1-3 ---
        while doc.TablesOfContents.Count > 0:
            doc.TablesOfContents(1).Delete()

        toc_range = None
        for i in range(1, doc.Paragraphs.Count + 1):
            txt = doc.Paragraphs(i).Range.Text.replace("\r", "").strip().upper()
            if txt == "TABLE OF CONTENTS":
                r = doc.Paragraphs(i).Range
                r.Collapse(0)  # end
                toc_range = r
                break

        if toc_range is not None:
            toc = doc.TablesOfContents.Add(
                Range=toc_range,
                UseHeadingStyles=True,
                UpperHeadingLevel=1,
                LowerHeadingLevel=3,
                UseHyperlinks=True,
                IncludePageNumbers=True,
                RightAlignPageNumbers=True,
            )
            toc.Update()
            print("TOC rebuilt with Heading 1–3.")

        # Update LOF/LOT and all fields
        doc.Fields.Update()
        try:
            for i in range(1, doc.TablesOfFigures.Count + 1):
                doc.TablesOfFigures(i).Update()
        except Exception:
            pass

        doc.Repaginate()
        doc.Save()
        doc.Close(False)
        print("Saved:", OUT)
    finally:
        word.Quit()


def main():
    # Close any open Word locks first is caller's responsibility
    normalize_headings_python()
    setup_page_numbers_and_toc()


if __name__ == "__main__":
    main()
