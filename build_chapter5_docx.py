"""Build Chapter 5 Word document for QR PC Check System."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "Chapter5_Figures"
DOC_PATH = ROOT / "CHAPTER_FIVE_Implementation_Coding_and_Testing.docx"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading_styled(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level <= 2 else 12, bold=True)
    return p


def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullet(doc, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    # clear default and set font
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0])
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True)
    return p


def add_figure(doc, filename, caption, width=5.6):
    path = FIG / filename
    if not path.exists():
        add_body(doc, f"[Screenshot placeholder: {filename}]", first_line_indent=False)
        add_caption(doc, caption)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        set_run_font(r, size=11, bold=True)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
        set_cell_shading(cell, "1F4E79")
        run.font.color.rgb = RGBColor(255, 255, 255)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()
    return table


def add_code(doc, code_text, title=None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_run_font(run, size=11, bold=True, italic=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text.strip("\n"))
    set_run_font(run, name="Consolas", size=8)
    # light background via shading on paragraph is complex; keep monospace
    return p


def add_mono_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run(lines)
    set_run_font(run, name="Consolas", size=9)
    return p


CONNECTION_CODE = r'''<?php
try {
    $pdo = new PDO(
        "mysql:host=localhost;dbname=computer_records;",
        'root',
        '',
        [PDO::ATTR_ERRMODE => PDO::ERRMODE_EXCEPTION]
    );
} catch (PDOException $e) {
    die('Connection failed: ' . $e->getMessage());
}
?>'''

LOGIN_CODE = r'''<?php
session_start();
include 'connection.php';

function authenticateUser($email, $password, $user_type, $pdo) {
    $query = "SELECT * FROM users WHERE email = :email AND user_type = :user_type";
    $stmt = $pdo->prepare($query);
    $stmt->bindParam(':email', $email);
    $stmt->bindParam(':user_type', $user_type);
    $stmt->execute();
    $row = $stmt->fetch(PDO::FETCH_ASSOC);
    if ($row && password_verify($password, $row['password'])) {
        return $row;
    }
    return false;
}

if ($_SERVER['REQUEST_METHOD'] === 'POST') {
    $user = authenticateUser($_POST['email'], $_POST['password'], $_POST['user_type'], $pdo);
    if ($user) {
        $_SESSION['user_type'] = $user['user_type'];
        $_SESSION['email'] = $user['email'];
        if ($user['user_type'] === 'Admin') {
            header('Location: admin-dashboard.php');
        } else {
            header('Location: user-dashboard.php');
        }
        exit();
    }
    $error = "Invalid user_type, username or password";
}
?>'''

RECORD_CODE = r'''<?php
require_once 'connection.php';

if (isset($_POST['submit'])) {
    $sn = $_POST['sn'];
    $model = $_POST['model'];
    $type = $_POST['type'];
    $owno = $_POST['owno'];
    $owname = $_POST['owname'];

    $sql_check = "SELECT COUNT(*) FROM computer_info WHERE sn = :sn";
    $stmt_check = $pdo->prepare($sql_check);
    $stmt_check->bindParam(':sn', $sn);
    $stmt_check->execute();

    if ($stmt_check->fetchColumn() > 0) {
        $Error = "Error: Computer's S/N already exists.";
    } else {
        $stmt = $pdo->prepare(
            "INSERT INTO computer_info (sn, model, type, owno, owname) VALUES (?,?,?,?,?)"
        );
        $stmt->execute([$sn, $model, $type, $owno, $owname]);
        $message = "Computer was recorded successfully!";
    }
}
?>'''

LOG_CODE = r'''<?php
require_once 'connection.php';

if ($_SERVER['REQUEST_METHOD'] == 'POST') {
    $stmt = $pdo->prepare(
        "INSERT INTO logs (sn, model, type, owno, owname, action, comment)
         VALUES (?,?,?,?,?,?,?)"
    );
    $stmt->execute([
        $_POST['sn'], $_POST['model'], $_POST['type'],
        $_POST['owno'], $_POST['owname'], $_POST['action'], $_POST['comment']
    ]);
    echo "Log recorded successfully for " . htmlspecialchars($_POST['sn']);
}
?>'''

INDEX_HTML = r'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Landing Page</title>
  <link rel="stylesheet"
        href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css">
</head>
<body>
  <div class="container">
    <h4>Welcome to the System for personal computers
        check using QR Codes</h4>
    <a href="login.php" class="btn btn-lg">Get Started</a>
  </div>
</body>
</html>'''


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CHAPTER 5: IMPLEMENTATION, CODING AND TESTING")
    set_run_font(r, size=16, bold=True)

    # 5.1
    add_heading_styled(doc, "5.1 Implementation and coding", 2)
    add_heading_styled(doc, "5.1.1 Introduction", 3)
    add_body(
        doc,
        "This chapter describes the implementation and testing of the System for Personal Computer "
        "Checks Using QR Codes developed for UTB Rubavu Campus. This phase focused on transforming "
        "the proposed design into a fully functional web-based application.",
    )
    add_body(
        doc,
        "The implementation phase explains how the system was developed using the selected "
        "technologies and tools. It covers frontend development, backend development with PHP and "
        "PDO, MySQL database integration, session-based authentication, QR code generation, and "
        "local deployment using XAMPP. The testing section explains how the system was evaluated to "
        "ensure that it operated correctly, securely, and efficiently according to the defined "
        "requirements.",
    )

    # 5.1.2
    add_heading_styled(doc, "5.1.2 Description of implementation Tools and technology", 3)
    add_body(
        doc,
        "The System for Personal Computer Checks Using QR Codes was developed using web technologies "
        "suitable for building a reliable campus gate-checking application. The system followed a "
        "three-tier client–server architecture, where the presentation layer (HTML, Bootstrap, and "
        "JavaScript) communicates with the application layer (PHP with PDO), which then interacts "
        "with the MySQL database. The technologies used in the development of the system are "
        "summarized in the table below.",
    )

    add_table(
        doc,
        ["Technology used", "Description"],
        [
            ["HTML5 (Frontend)", "Used to structure web pages including the landing page, login page, admin dashboard, gate-officer dashboard, registration forms, and report pages."],
            ["CSS3 (Frontend)", "Used to style and format the web pages to improve appearance, layout, and readability."],
            ["Bootstrap 4", "Used to create responsive and consistent interfaces across devices and browsers."],
            ["JavaScript (Frontend)", "Used to enhance client-side interaction, form behaviour, and dynamic page features."],
            ["PHP with PDO (Backend)", "Used as the server-side language to process requests, enforce business rules, and interact securely with the database."],
            ["MySQL", "Used as the relational database management system to store users, computer records, and activity logs."],
            ["TCPDF / QR libraries", "Used to generate printable QR codes linked to registered computer serial numbers."],
            ["XAMPP", "Used as the local development and hosting environment (Apache, MySQL, and PHP)."],
            ["phpMyAdmin", "Used to create, import, and manage the computer_records database."],
            ["Visual Studio Code / Cursor", "Used as the primary code editor for development and debugging."],
        ],
        caption="Table 6: Technology used",
    )
    add_body(
        doc,
        "These technologies were selected because they provide flexibility, cross-platform "
        "compatibility, secure database access through prepared statements, and an efficient "
        "development workflow for academic and institutional web systems.",
        first_line_indent=False,
    )

    # 5.1.3 Screenshots
    add_heading_styled(doc, "5.1.3 Screen Shots", 3)

    add_figure(doc, "Figure10_Home_Page.png", "Figure 10: Home Page Interface")
    add_body(
        doc,
        "This figure presents the Home Page of the System for Personal Computer Checks Using QR "
        "Codes. It serves as the main entry point for authorized campus users and provides a clear "
        "welcome message together with a “Get Started” button that directs users to the login page.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure11_Login_Page.png", "Figure 11: Login Page")
    add_body(
        doc,
        "This figure illustrates the Login page, where users select their role (Admin or Gate "
        "Officer), enter their email address and password, and authenticate into the system. Role "
        "selection ensures that each user is redirected to the appropriate dashboard after a "
        "successful login.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure12_Admin_Dashboard.png", "Figure 12: Admin Dashboard")
    add_body(
        doc,
        "This figure presents the Admin Dashboard. From this interface, administrators can manage "
        "system users, including adding, viewing, updating, and deleting accounts. The dashboard "
        "acts as the central control panel for system administration at UTB Rubavu Campus.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure13_View_Users.png", "Figure 13: View Users Page")
    add_body(
        doc,
        "This figure shows the View Users page, which lists registered system accounts with their "
        "roles, identification numbers, names, and email addresses. It enables administrators to "
        "monitor who has access to the platform and to maintain accurate user records.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure14_Add_Users.png", "Figure 14: Add Users Page")
    add_body(
        doc,
        "This figure illustrates the Add Users interface. Administrators can register new Admin or "
        "Gate Officer accounts by capturing user type, national/staff identification number, full "
        "name, and email. Passwords are stored in hashed form for security.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure15_User_Dashboard.png", "Figure 15: Gate Officer Dashboard")
    add_body(
        doc,
        "This figure presents the Gate Officer (Guest) Dashboard. From this panel, gate staff can "
        "access computer registration, laptop viewing, QR-related operations, password change, and "
        "reporting features used during daily entry and exit checks.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure16_Record_Computers.png", "Figure 16: Record Computers Page")
    add_body(
        doc,
        "This figure shows the Record Computers page used to register a personal computer at the "
        "gate. Gate officers capture the serial number, model, owner type, owner number, and owner "
        "name. After successful registration, a QR code can be generated for the device.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure17_View_Laptops.png", "Figure 17: View Laptops / Registered Computers Page")
    add_body(
        doc,
        "This figure illustrates the View Laptops page, which displays registered computer records "
        "stored in the database. Gate officers can review device details and use the information "
        "during verification and QR-code generation activities.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure18_Report_Page.png", "Figure 18: Reports Module")
    add_body(
        doc,
        "This figure presents the Reports module. It enables authorized users to review logged "
        "check-in and check-out activity and to generate printable reports for accountability and "
        "management purposes at the campus gate.",
        first_line_indent=False,
    )

    add_figure(doc, "Figure19_Change_Password.png", "Figure 19: Change Password Page")
    add_body(
        doc,
        "This figure shows the Change Password page, which allows authenticated users to update "
        "their login credentials. This feature supports account security and controlled access to "
        "system modules.",
        first_line_indent=False,
    )

    # 5.1.4 Source Codes
    add_heading_styled(doc, "5.1.4 Source Codes", 3)
    add_body(
        doc,
        "The System for Personal Computer Checks Using QR Codes was implemented as a PHP web "
        "application hosted under the Apache document root. The main project folder contains "
        "presentation pages, backend scripts, database connection logic, and libraries used for QR "
        "code generation.",
        first_line_indent=False,
    )

    p = doc.add_paragraph()
    run = p.add_run("PROJECT FOLDER STRUCTURE")
    set_run_font(run, size=12, bold=True)

    add_body(doc, "The application folder structure is summarized below:", first_line_indent=False)
    add_mono_block(
        doc,
        "QR/\n"
        "├── index.php                 # Landing / home page\n"
        "├── login.php                 # Authentication page\n"
        "├── logout.php                # Session termination\n"
        "├── connection.php            # MySQL PDO connection\n"
        "├── admin-dashboard.php       # Administrator home\n"
        "├── user-dashboard.php        # Gate officer home\n"
        "├── add-users.php             # Create system users\n"
        "├── view-users.php            # List / manage users\n"
        "├── update-users.php          # Edit user accounts\n"
        "├── delete-users.php          # Remove user accounts\n"
        "├── record-computers.php      # Register devices\n"
        "├── view-laptops.php          # List registered devices\n"
        "├── getBarcode.php            # QR display helper\n"
        "├── log_form.php / submit_log.php  # Check-in/out logging\n"
        "├── report.php / generate_report.php\n"
        "├── change-password.php\n"
        "├── styles.css\n"
        "├── img/                      # Interface images\n"
        "├── icons/                    # Icon assets\n"
        "├── tcpdf/                    # PDF / barcode library\n"
        "└── phpqr/                    # QR generation support",
    )

    add_body(
        doc,
        "Sample database connection file (connection.php):",
        first_line_indent=False,
    )
    add_code(doc, CONNECTION_CODE)

    add_body(
        doc,
        "Sample authentication logic (login.php):",
        first_line_indent=False,
    )
    add_code(doc, LOGIN_CODE)

    add_body(
        doc,
        "Sample computer registration logic (record-computers.php):",
        first_line_indent=False,
    )
    add_code(doc, RECORD_CODE)

    add_body(
        doc,
        "Sample check-in / check-out logging logic (submit_log.php):",
        first_line_indent=False,
    )
    add_code(doc, LOG_CODE)

    p = doc.add_paragraph()
    run = p.add_run("FRONTEND SAMPLE")
    set_run_font(run, size=12, bold=True)

    add_body(
        doc,
        "The frontend of the System for Personal Computer Checks Using QR Codes was developed using "
        "HTML, CSS, Bootstrap, and JavaScript. A simplified landing-page structure is shown below.",
        first_line_indent=False,
    )
    add_code(doc, INDEX_HTML)

    # 5.2 Testing
    add_heading_styled(doc, "5.2 Testing", 2)
    add_heading_styled(doc, "5.2.1 Introduction", 3)
    add_body(
        doc,
        "System testing is the process of evaluating a complete and integrated system to ensure that "
        "it meets the specified requirements. Testing involves executing the system to identify "
        "errors, defects, or inconsistencies and verifying that all components function as expected.",
    )
    add_body(
        doc,
        "The System for Personal Computer Checks Using QR Codes underwent several types of testing "
        "to ensure reliability, usability, security, and performance before final demonstration and "
        "local deployment at UTB Rubavu Campus.",
    )

    add_heading_styled(doc, "5.2.2 Objective of Testing", 3)
    add_body(doc, "The primary objectives of system testing in this project were:", first_line_indent=False)
    for i, item in enumerate([
        "To identify and correct errors during development.",
        "To ensure that all functional requirements are implemented correctly.",
        "To verify that the system meets administrator and gate-officer expectations.",
        "To confirm that the system operates securely and efficiently.",
        "To ensure compatibility across different browsers and devices.",
        "To validate that the system performs according to the System Requirement Specification (SRS).",
    ], 1):
        add_bullet(doc, f"{i}.\t{item}")
    add_body(
        doc,
        "Testing ensured that the platform supports accurate device registration, QR-based "
        "verification, secure authentication, and reliable activity logging at the campus gate.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "5.2.3 Unit testing outputs", 3)
    add_body(
        doc,
        "Unit testing was conducted to verify individual components of the system. PHP scripts and "
        "database operations were tested independently to ensure correct responses. Unit testing included:",
    )
    for i, item in enumerate([
        "Testing the PDO database connection to confirm successful access to the computer_records database.",
        "Testing user authentication with valid and invalid credentials for Admin and Gate Officer roles.",
        "Testing INSERT operations when registering a new computer record.",
        "Testing duplicate serial-number prevention during computer registration.",
        "Testing QR code generation for a registered serial number.",
        "Testing INSERT operations into the logs table for check-in and check-out actions.",
        "Testing password hashing and password update operations.",
    ], 1):
        add_bullet(doc, f"{i}.\t{item}")
    add_body(
        doc,
        "Each function was tested independently before integrating with other components.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "5.2.4 Validation testing outputs", 3)
    add_body(
        doc,
        "Validation testing ensured that the system behaved according to user expectations. Input "
        "fields were validated to prevent incorrect data submission. This included:",
    )
    for i, item in enumerate([
        "Ensuring required fields are completed before form submission on login, user registration, and computer registration forms.",
        "Validating email format on the login and user creation forms.",
        "Preventing unauthorized access to admin and gate-officer dashboards through session checks.",
        "Displaying appropriate error messages for invalid login credentials.",
        "Rejecting duplicate computer serial numbers during registration.",
    ], 1):
        add_bullet(doc, f"{i}.\t{item}")
    add_body(
        doc,
        "Validation confirmed that the system performs intended operations correctly and reduces "
        "the risk of incomplete or inconsistent data entry.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "5.2.5 Integration Testing Outputs", 3)
    add_body(
        doc,
        "Integration testing was conducted to verify that frontend pages, PHP backend scripts, and "
        "the MySQL database work together properly. This included:",
    )
    for i, item in enumerate([
        "Testing login form submission and redirection to the correct dashboard based on user role.",
        "Testing computer registration forms and immediate storage of records in the Computer_info table.",
        "Testing QR code generation linked to newly registered device serial numbers.",
        "Testing check-in/check-out logging and storage of actions in the Logs table.",
        "Testing report generation from stored log and device data.",
        "Testing admin user-management operations (add, view, update, and delete) against the Users table.",
    ], 1):
        add_bullet(doc, f"{i}.\t{item}")
    add_body(
        doc,
        "This testing ensured smooth communication between the presentation, application, and "
        "database layers of the system.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "5.2.6 Functional and system testing Results", 3)
    add_body(
        doc,
        "Functional testing verified that each feature operates according to system requirements. "
        "All core functionalities were tested, including:",
    )
    for i, item in enumerate([
        "Accessing the home page and navigating to login",
        "Admin and Gate Officer authentication",
        "Adding, viewing, updating, and deleting users",
        "Recording personal computers at entry",
        "Viewing registered laptop/computer records",
        "Generating and displaying QR codes",
        "Logging check-in and check-out actions",
        "Generating activity reports",
        "Changing user passwords",
        "Logging out and terminating sessions securely",
    ], 1):
        add_bullet(doc, f"{i}.\t{item}")
    add_body(
        doc,
        "The system responded correctly to user inputs and displayed accurate results. No major "
        "functional defects were identified after final testing.",
        first_line_indent=False,
    )

    add_heading_styled(doc, "5.2.7 Acceptance Testing Report", 3)
    add_body(
        doc,
        "Acceptance testing was conducted to determine whether the system meets the needs of end "
        "users at UTB Rubavu Campus. The platform was tested by simulating real gate operations such "
        "as registering a device on entry, generating a QR code, and recording verification actions "
        "on exit.",
    )
    add_body(
        doc,
        "Administrators tested dashboard functionality, including adding, editing, and deleting user "
        "accounts. Gate officers tested computer registration, QR-related operations, and reporting. "
        "Feedback confirmed that the system is usable, responsive under the local XAMPP environment, "
        "and meets the intended objectives of improving personal computer checks through QR codes "
        "at the campus gate.",
    )

    doc.save(DOC_PATH)
    print("Saved:", DOC_PATH)


if __name__ == "__main__":
    build()
